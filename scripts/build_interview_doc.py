from pathlib import Path
import re
import argparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('第 ')
    set_font(run, size=9, color='6B7280')
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)
    run = paragraph.add_run(' 页')
    set_font(run, size=9, color='6B7280')


def add_text(paragraph, text):
    # Lightweight inline emphasis for the source markdown.
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith('**') and part.endswith('**')
        value = part[2:-2] if bold else part
        run = paragraph.add_run(value)
        set_font(run, bold=bold)


def paragraph(doc, text='', style=None, bold=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_font(r, bold=bold)
    return p


def build(source_path=None, output_path=None):
    src = Path(source_path) if source_path else ROOT / 'docs' / '中兴面试半天冲刺准备.md'
    out = Path(output_path) if output_path else ROOT / '中兴面试半天冲刺准备.docx'
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ('Heading 1', 16, '2E74B5', 18, 10),
        ('Heading 2', 13, '2E74B5', 14, 7),
        ('Heading 3', 12, '1F4D78', 10, 5),
    ]:
        st = styles[name]
        st.font.name = 'Calibri'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    r = header.add_run('中兴面试半天冲刺准备 | 杨立昌')
    set_font(r, size=9, color='6B7280')
    add_page_number(section.footer.paragraphs[0])

    lines = src.read_text(encoding='utf-8').splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line[2:])
            set_font(r, size=22, bold=True, color='0B2545')
            first_title = False
            i += 1
            continue
        if line.startswith('## '):
            p = doc.add_paragraph(line[3:], style='Heading 1')
            for r in p.runs:
                set_font(r, size=16, bold=True, color='2E74B5')
            i += 1
            continue
        if line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='Heading 2')
            for r in p.runs:
                set_font(r, size=13, bold=True, color='2E74B5')
            i += 1
            continue
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(8)
            add_text(p, line[2:])
            p._p.get_or_add_pPr().append(OxmlElement('w:shd'))
            p._p.pPr[-1].set(qn('w:fill'), 'F4F6F9')
            i += 1
            continue
        if line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                cells = [c.strip() for c in lines[i].strip('|').split('|')]
                if not all(re.fullmatch(r':?-{3,}:?', c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.style = 'Table Grid'
                widths = [2700, 6660] if cols == 2 else [3120] * cols
                if sum(widths) != 9360:
                    widths[-1] += 9360 - sum(widths)
                set_table_width(table, widths)
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        cell = table.cell(ri, ci)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        add_text(p, row[ci] if ci < len(row) else '')
                        if ri == 0:
                            shade(cell, 'E8EEF5')
                            for r in p.runs:
                                r.bold = True
                doc.add_paragraph().paragraph_format.space_after = Pt(0)
            continue
        if re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            add_text(p, re.sub(r'^\d+\. ', '', line))
            i += 1
            continue
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            add_text(p, line[2:])
            i += 1
            continue
        p = doc.add_paragraph()
        add_text(p, line)
        i += 1

    doc.core_properties.title = '中兴面试半天冲刺准备'
    doc.core_properties.subject = '系统底层与驱动开发方向面试准备'
    doc.core_properties.author = '杨立昌'
    doc.save(out)
    print(out)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--source')
    parser.add_argument('--output')
    args = parser.parse_args()
    build(args.source, args.output)
